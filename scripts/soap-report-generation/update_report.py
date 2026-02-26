#!/usr/bin/env python3
"""
Update Word delivery report by replacing Summary section with content from summary.md
"""
import logging
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)-7s %(message)s")
log = logging.getLogger(__name__)


"""Parse and validate date string in YYYY-MM-DD or YYYY_MM_DD format"""
def parse_date(date_str):
  try:
    dt = datetime.strptime(date_str, '%Y-%m-%d')
    return dt.year, dt.month, dt.day
  except ValueError:
    try:
      dt = datetime.strptime(date_str, '%Y_%m_%d')
      return dt.year, dt.month, dt.day
    except ValueError:
      raise ValueError(f"Invalid date format: {date_str}. Expected YYYY-MM-DD")


"""Find the paragraph indices where Summary section starts and ends"""
def find_summary_section_bounds(doc):
  summary_start = None
  summary_end = None

  for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    if summary_start is None:
      if re.match(r'^1\.?\s*Summary$|^Summary$|^1\tSummary$', text):
        summary_start = i
        log.debug(f"Found Summary section at paragraph {i}: '{text}'")
        continue

    if summary_start is not None and summary_end is None:
      if re.match(r'^2[.\s\t]', text):
        summary_end = i
        log.debug(f"Summary section ends at paragraph {i}")
        break

  if summary_start is not None and summary_end is None:
    summary_end = len(doc.paragraphs)
    log.warning("No section 2 found, using end of document")

  return summary_start, summary_end


"""Replace content between Summary heading and next section with new content"""
def replace_summary_content(doc, summary_docx):
  summary_start, summary_end = find_summary_section_bounds(doc)

  if summary_start is None:
    raise ValueError("Could not find Summary section in document")

  # Remove paragraphs between heading and next section (reverse order)
  removed_count = 0
  for i in range(summary_end - 1, summary_start, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)
    removed_count += 1

  log.debug(f"Removed paragraphs from index {summary_start + 1} to {summary_end - 1}")

  # Delete second table if exists (part of old summary structure)
  if len(doc.tables) > 1:
    table = doc.tables[1]
    table._element.getparent().remove(table._element)
    log.debug("Removed second table from document")

  # Insert new content after summary heading
  for i in range(len(summary_docx.paragraphs)):
    if i == 0:
      doc.add_heading(summary_docx.paragraphs[i].text, level=2)
    else:
      doc.add_paragraph(summary_docx.paragraphs[i].text, style='List Bullet')

  log.debug(f"Inserted new summary content")


"""Update date references on the first page"""
def update_dates_on_first_page(doc, old_date, new_date):
  old_year, old_month, old_day = parse_date(old_date)
  new_year, new_month, new_day = parse_date(new_date)

  old_tag = f"{old_year}/{old_month:02d}/delivery-{old_year}-{old_month:02d}-{old_day:02d}"
  new_tag = f"{new_year}/{new_month:02d}/delivery-{new_year}-{new_month:02d}-{new_day:02d}"

  for para in doc.paragraphs[:20]:
    if f"Tag {old_tag}" in para.text:
      para.text = f"Tag {new_tag}"
      for run in para.runs:
        run.bold = True
        run.font.size = Pt(18)
    elif para.text.startswith("(including deliveries after"):
      para.text = f"(including deliveries after {old_tag})"

  log.debug("Updated date references on the first page")


"""Add AI generation note if paragraph 2 is empty"""
def add_ai_note_if_missing(doc, ai_note):
  if len(doc.paragraphs) < 3:
    log.warning("Document has fewer than 3 paragraphs, cannot add AI note")
    return

  target = doc.paragraphs[2]
  if not target.text.strip():
    target.text = ai_note
    target.alignment = 1
    for run in target.runs:
      run.font.size = Pt(8)
      run.font.name = 'Lucida Console'
    log.debug("Added AI generation note")
  else:
    log.debug("Paragraph 2 not empty, skipping AI note")


def main():
  if len(sys.argv) != 3:
    log.error("Invalid arguments. Expected old_date and new_date in YYYY-MM-DD format. Usage: update_report.py <old_date> <new_date>")
    sys.exit(1)

  old_date = sys.argv[1]
  new_date = sys.argv[2]

  # Validate dates early
  try:
    parse_date(old_date)
    parse_date(new_date)
  except ValueError as e:
    log.error("Date validation failed")
    sys.exit(1)

  # Use relative path
  tmp_dir = Path(__file__).parent.parent / "tmp"

  old_date_file = old_date.replace('-', '_')
  new_date_file = new_date.replace('-', '_')

  summary_file = tmp_dir / "summary.docx"
  src_doc = tmp_dir / f"SOAP_API_delivery_report_{old_date_file}.docx"
  dest_doc = tmp_dir / f"SOAP_API_delivery_report_{new_date_file}.docx"

  log.info("Step 1: Verify summary.docx exists")
  if not summary_file.exists():
    log.error("%s not found", str(summary_file))
    sys.exit(1)

  log.info("Step 2: Load summary.docx")
  summary_docx = Document(summary_file)

  log.info("Step 3: Verify source document exists")
  if not src_doc.exists():
    log.error("%s not found", str(src_doc))
    sys.exit(1)

  log.info("Step 4: Load source document")
  try:
    doc = Document(src_doc)
  except Exception as e:
    log.error("Failed to load source document: %s", e)
    sys.exit(1)

  log.info("Step 5: Update date references on first page")
  update_dates_on_first_page(doc, old_date, new_date)

  log.info("Step 6: Add AI generation note if missing")
  add_ai_note_if_missing(doc, "Note: This report was generated by an AI agent.")

  log.info("Step 7: Replace Summary section content")
  try:
    replace_summary_content(doc, summary_docx)
  except ValueError as e:
    log.error("Failed to replace summary: %s", e)
    sys.exit(1)

  log.info("Step 8: Save updated document")
  try:
    doc.save(str(dest_doc))
  except Exception as e:
    log.error("Failed to save document: %s", e)
    sys.exit(1)

  log.info("SUCCESS: Updated report saved to %s", str(dest_doc))


if __name__ == "__main__":
  main()
